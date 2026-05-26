from django.shortcuts import render, redirect
from django.contrib.auth.mixins import LoginRequiredMixin
from django.views.generic import ListView, CreateView, UpdateView, DeleteView
from django.urls import reverse_lazy
from django.contrib import messages
from django.http import JsonResponse
from .models import Equipment, EquipmentCategory
from apps.users.views import AdminRequiredMixin

class EquipmentListView(LoginRequiredMixin, ListView):
    model = Equipment
    template_name = 'inventory/equipment_list.html'
    context_object_name = 'equipment_list'
    paginate_by = 10

    def get(self, request, *args, **kwargs):
        # Clamp page number to valid range
        try:
            page = int(request.GET.get('page', 1))
            if page < 1:
                from django.http import HttpResponseRedirect
                params = request.GET.copy()
                params['page'] = 1
                return HttpResponseRedirect(f"?{params.urlencode()}")
        except (ValueError, TypeError):
            pass
        return super().get(request, *args, **kwargs)

    def get_queryset(self):
        queryset = super().get_queryset()
        
        # Search
        search = self.request.GET.get('search')
        if search:
            queryset = queryset.filter(name__icontains=search) | queryset.filter(inventory_code__icontains=search)
        
        # Filter by status
        status = self.request.GET.get('status')
        if status:
            queryset = queryset.filter(status=status)
        
        # Filter by main category or subcategory
        main_category = self.request.GET.get('main_category')
        category = self.request.GET.get('category')
        
        if category:
            # If subcategory is selected, filter by subcategory
            queryset = queryset.filter(category_id=category)
        elif main_category:
            # If only main category is selected, filter by main category and all its children
            main_cat = EquipmentCategory.objects.filter(id=main_category).first()
            if main_cat:
                # Get all subcategories of the main category
                subcategory_ids = list(main_cat.get_children().values_list('id', flat=True))
                subcategory_ids.append(int(main_category))  # Include main category itself
                queryset = queryset.filter(category_id__in=subcategory_ids)
        
        return queryset
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['categories'] = EquipmentCategory.objects.all()
        context['root_categories'] = EquipmentCategory.get_root_categories()
        context['status_choices'] = Equipment.STATUS_CHOICES
        return context

class EquipmentCreateView(AdminRequiredMixin, CreateView):
    model = Equipment
    template_name = 'inventory/equipment_form.html'
    fields = ['name', 'inventory_code', 'category', 'specifications', 'current_user', 
              'location', 'status', 'purchase_date', 'warranty_until']
    success_url = reverse_lazy('inventory:equipment_list')
    
    def form_valid(self, form):
        messages.success(self.request, 'Peralatan berhasil ditambahkan!')
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['categories'] = EquipmentCategory.objects.all()
        context['root_categories'] = EquipmentCategory.get_root_categories()
        return context

class EquipmentUpdateView(AdminRequiredMixin, UpdateView):
    model = Equipment
    template_name = 'inventory/equipment_form.html'
    fields = ['name', 'inventory_code', 'category', 'specifications', 'current_user', 
              'location', 'status', 'purchase_date', 'warranty_until']
    success_url = reverse_lazy('inventory:equipment_list')
    
    def form_valid(self, form):
        messages.success(self.request, 'Peralatan berhasil diupdate!')
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['categories'] = EquipmentCategory.objects.all()
        context['root_categories'] = EquipmentCategory.get_root_categories()
        return context

class EquipmentDeleteView(AdminRequiredMixin, DeleteView):
    model = Equipment
    template_name = 'inventory/equipment_confirm_delete.html'
    success_url = reverse_lazy('inventory:equipment_list')
    
    def delete(self, request, *args, **kwargs):
        messages.success(request, 'Peralatan berhasil dihapus!')
        return super().delete(request, *args, **kwargs)

# Category Management Views
class CategoryListView(AdminRequiredMixin, ListView):
    model = EquipmentCategory
    template_name = 'inventory/category_list.html'
    context_object_name = 'categories'
    paginate_by = 20

class CategoryCreateView(AdminRequiredMixin, CreateView):
    model = EquipmentCategory
    template_name = 'inventory/category_form.html'
    fields = ['parent', 'name', 'description']
    success_url = reverse_lazy('inventory:category_list')
    
    def form_valid(self, form):
        messages.success(self.request, 'Kategori berhasil ditambahkan!')
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        context['root_categories'] = EquipmentCategory.get_root_categories()
        return context

class CategoryUpdateView(AdminRequiredMixin, UpdateView):
    model = EquipmentCategory
    template_name = 'inventory/category_form.html'
    fields = ['parent', 'name', 'description']
    success_url = reverse_lazy('inventory:category_list')
    
    def form_valid(self, form):
        messages.success(self.request, 'Kategori berhasil diupdate!')
        return super().form_valid(form)
    
    def get_context_data(self, **kwargs):
        context = super().get_context_data(**kwargs)
        # Exclude current category from parent options to prevent circular reference
        context['root_categories'] = EquipmentCategory.get_root_categories().exclude(id=self.object.id)
        return context

class CategoryDeleteView(AdminRequiredMixin, DeleteView):
    model = EquipmentCategory
    template_name = 'inventory/category_confirm_delete.html'
    success_url = reverse_lazy('inventory:category_list')
    
    def delete(self, request, *args, **kwargs):
        messages.success(request, 'Kategori berhasil dihapus!')
        return super().delete(request, *args, **kwargs)


from django.http import HttpResponse
from django.contrib.auth.decorators import login_required
from django.views.decorators.http import require_POST
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill
from datetime import datetime
import json
import qrcode
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT

@login_required
def get_subcategories(request):
    """AJAX view untuk mendapatkan sub-kategori berdasarkan parent"""
    parent_id = request.GET.get('parent_id')
    if parent_id:
        subcategories = EquipmentCategory.objects.filter(parent_id=parent_id)
        data = [{'id': cat.id, 'name': cat.name} for cat in subcategories]
        return JsonResponse({'subcategories': data})
    return JsonResponse({'subcategories': []})

def generate_qr_code(request, equipment_id):
    """Generate and download QR code image for equipment"""
    try:
        equipment = Equipment.objects.get(id=equipment_id)
        
        # Create readable QR data (not JSON format)
        qr_text = f"""IT HUB INVENTORY SYSTEM
LAPORAN DETAIL PERALATAN

Nama Peralatan: {equipment.name}
Kode Inventaris: {equipment.inventory_code}
Kategori: {equipment.category.name}
Status: {equipment.get_status_display()}
Pengguna Saat Ini: {equipment.current_user or 'Tidak ada'}
Lokasi: {equipment.location or 'Tidak diketahui'}
Spesifikasi: {equipment.specifications or 'Tidak ada spesifikasi'}
Tanggal Pembelian: {equipment.purchase_date.strftime('%d/%m/%Y') if equipment.purchase_date else 'Tidak diketahui'}
Garansi Sampai: {equipment.warranty_until.strftime('%d/%m/%Y') if equipment.warranty_until else 'Tidak diketahui'}

Dibuat: {equipment.created_at.strftime('%d/%m/%Y %H:%M')}
Diupdate: {equipment.updated_at.strftime('%d/%m/%Y %H:%M')}

Laporan dibuat pada: {datetime.now().strftime('%d/%m/%Y %H:%M:%S')}
Dibuat oleh: {request.user.username if request.user.is_authenticated else 'Anonymous'}"""
        
        # Create QR code
        qr = qrcode.QRCode(
            version=1,  # Controls the size of the QR Code
            error_correction=qrcode.constants.ERROR_CORRECT_L,  # About 7% or less errors can be corrected
            box_size=10,  # Controls how many pixels each "box" of the QR code is
            border=4,  # Controls how many boxes thick the border should be
        )
        
        # Add data to QR code
        qr.add_data(qr_text)
        qr.make(fit=True)
        
        # Create QR code image
        qr_img = qr.make_image(fill_color="black", back_color="white")
        
        # Save image to BytesIO buffer
        buffer = BytesIO()
        qr_img.save(buffer, format='PNG')
        buffer.seek(0)
        
        # Create HTTP response with image
        response = HttpResponse(buffer.getvalue(), content_type='image/png')
        filename = f'QR-{equipment.inventory_code}-{equipment.name.replace(" ", "_").replace("/", "_")}.png'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        
        return response
        
    except Equipment.DoesNotExist:
        return HttpResponse('Equipment not found', status=404)
    except Exception as e:
        # Fallback to text file if QR generation fails
        equipment = Equipment.objects.get(id=equipment_id)
        fallback_text = f"""IT HUB INVENTORY - QR FALLBACK

Item: {equipment.name}
Code: {equipment.inventory_code}
Category: {equipment.category.name}
User: {equipment.current_user or 'Tidak ada'}
Location: {equipment.location or 'Tidak diketahui'}
Status: {equipment.get_status_display()}

Error generating QR image: {str(e)}
Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
"""
        
        response = HttpResponse(fallback_text, content_type='text/plain')
        filename = f'QR-FALLBACK-{equipment.inventory_code}.txt'
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        return response



def word_document(request, equipment_id):
    return HttpResponse("Word function works...")


def _selected_equipment_queryset(request):
    selected_ids = request.POST.getlist('selected_equipment')
    return Equipment.objects.select_related('category').filter(id__in=selected_ids).order_by('name')


def _build_equipment_excel_response(equipment_list, filename_prefix='Inventaris_Terpilih'):
    wb = Workbook()
    ws = wb.active
    ws.title = "Inventaris Peralatan"

    header_fill = PatternFill(start_color="152231", end_color="152231", fill_type="solid")
    header_font = Font(bold=True, color="FFFFFF")
    header_alignment = Alignment(horizontal="center", vertical="center")

    headers = ['No', 'Nama', 'Kode Inventaris', 'Kategori', 'Spesifikasi', 'Pengguna', 'Lokasi', 'Status', 'Tanggal Pembelian', 'Garansi Sampai']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=header)
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = header_alignment

    for row, equipment in enumerate(equipment_list, 2):
        ws.cell(row=row, column=1, value=row - 1)
        ws.cell(row=row, column=2, value=equipment.name)
        ws.cell(row=row, column=3, value=equipment.inventory_code)
        ws.cell(row=row, column=4, value=equipment.category.full_path)
        ws.cell(row=row, column=5, value=equipment.specifications or '-')
        ws.cell(row=row, column=6, value=equipment.current_user or '-')
        ws.cell(row=row, column=7, value=equipment.location or '-')
        ws.cell(row=row, column=8, value=equipment.get_status_display())
        ws.cell(row=row, column=9, value=equipment.purchase_date.strftime('%d/%m/%Y') if equipment.purchase_date else '-')
        ws.cell(row=row, column=10, value=equipment.warranty_until.strftime('%d/%m/%Y') if equipment.warranty_until else '-')

    widths = [5, 28, 18, 24, 42, 22, 22, 15, 18, 18]
    for index, width in enumerate(widths, 1):
        ws.column_dimensions[chr(64 + index)].width = width

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    filename = f'{filename_prefix}_{datetime.now().strftime("%Y%m%d_%H%M%S")}.xlsx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    wb.save(response)
    return response


def _build_equipment_qr_word_response(request, equipment_list):
    document = Document()
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.45)
    section.right_margin = Inches(0.45)

    local_ip = None
    try:
        from apps.inventory.qr_views import get_local_ip
        local_ip = get_local_ip()
    except Exception:
        local_ip = request.get_host().split(':')[0]

    table = document.add_table(rows=0, cols=3)
    table.autofit = False
    for column in table.columns:
        for cell in column.cells:
            cell.width = Inches(3.05)

    for start in range(0, len(equipment_list), 3):
        row = table.add_row()
        row.height = Inches(2.45)
        for column_index, equipment in enumerate(equipment_list[start:start + 3]):
            cell = row.cells[column_index]
            cell.width = Inches(3.05)

            item_url = f"http://{local_ip}:9000/item/{equipment.inventory_code}/"
            qr = qrcode.QRCode(
                version=1,
                error_correction=qrcode.constants.ERROR_CORRECT_M,
                box_size=10,
                border=2,
            )
            qr.add_data(item_url)
            qr.make(fit=True)
            qr_img = qr.make_image(fill_color="black", back_color="white")

            image_buffer = BytesIO()
            qr_img.save(image_buffer, format='PNG')
            image_buffer.seek(0)

            qr_paragraph = cell.paragraphs[0]
            qr_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            qr_paragraph.add_run().add_picture(image_buffer, width=Inches(1.75))

            code_paragraph = cell.add_paragraph()
            code_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            code_run = code_paragraph.add_run(equipment.inventory_code)
            code_run.font.name = 'Times New Roman'
            code_run.font.size = Pt(10)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        buffer.getvalue(),
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )
    filename = f'QR_Inventaris_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx'
    response['Content-Disposition'] = f'attachment; filename="{filename}"'
    return response


@require_POST
@login_required
def bulk_equipment_action(request):
    if not request.user.is_staff:
        messages.error(request, 'Anda tidak memiliki akses untuk aksi massal inventaris.')
        return redirect('inventory:equipment_list')

    action = request.POST.get('bulk_action')
    equipment_list = list(_selected_equipment_queryset(request))

    if not equipment_list:
        messages.warning(request, 'Pilih minimal satu inventaris dulu.')
        return redirect('inventory:equipment_list')

    if action == 'delete':
        count = len(equipment_list)
        Equipment.objects.filter(id__in=[equipment.id for equipment in equipment_list]).delete()
        messages.success(request, f'{count} inventaris berhasil dihapus.')
        return redirect('inventory:equipment_list')

    if action == 'export_excel':
        return _build_equipment_excel_response(equipment_list)

    if action == 'export_qr_word':
        return _build_equipment_qr_word_response(request, equipment_list)

    messages.error(request, 'Aksi massal tidak valid.')
    return redirect('inventory:equipment_list')

@login_required
def export_equipment_excel(request):
    """Export inventaris peralatan ke Excel"""
    if not request.user.is_staff:
        messages.error(request, 'Anda tidak memiliki akses untuk export data!')
        return redirect('inventory:equipment_list')
    
    equipment_list = Equipment.objects.select_related('category').all().order_by('-created_at')
    return _build_equipment_excel_response(equipment_list, filename_prefix='Inventaris_Peralatan')
